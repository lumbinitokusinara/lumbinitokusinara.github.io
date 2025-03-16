import os
from docx2python import docx2python
from docx import Document
#import docx2txt
from bs4 import BeautifulSoup

input_folder = 'docs'
output_folder = 'docs/html'

pages_info = [
    ("Introduction.txt", "Introduction-.docx", ""),
    ("page1.txt", "Section 1-Birth.docx", "1."),
    #("page1.txt", "1 - Birth of the Buddha Gotama.docx", "1."),
    ("page2.txt", "Section 2- Childhood and Family of the Buddha.docx", "2."),
    ("page3.txt", "Section 3-Asetic life.docx", "3."),
    ("page4.txt", "Section 4 -Enlightenment.docx", "4."),
    ("page4_1.txt", "Section 4-1 Four Assembly.docx", "4-1."),
    ("page5.txt", "Section 5 DF- Greatness of The Buddha.docx", "5."),
    ("page6.txt", "Section 6 DF - Titles.docx", "6."),
    ("page7.txt", "Section 7 DF- As a  Teacher.docx", "7."),
    ("page8.txt", "8 Section-Praise and B.docx", "8."),
    ("page9.txt", "Section 9- DF  Past lives-.docx", "9."),
   
    ("page11_1.txt","Section 11 -1 DF- Advice to the Sangha.docx", "11.1"),
    ("page11_2.txt","Section 11-2  Advice to Lay Followers.docx", "11.2")
    
]
# Ensure the output folder exists
os.makedirs(output_folder, exist_ok=True)

filename = os.path.join(input_folder, "Section 4 -Enlightenment.docx")

# doc = docx.Document(filename)
html_content = docx2python(filename, html = True)
print(html_content.text)

#with docx2python(filename, html = True) as html_content:
#    print(html_content.text)
    
soup = BeautifulSoup(html_content.text, "html.parser")

# Find and remove all <span> tags with a style attribute
for span in soup.find_all("span", attrs={"style": True}):
    span.unwrap()  # Removes the tag but keeps its content

output_path = os.path.join(output_folder,'file4.txt')
with open(output_path, "w", encoding="utf-8") as html_file:
    html_file.write(str(soup))


# Generate the pages
for ofile, docfile, section_num in pages_info:
    outputfile = os.path.join(output_folder, ofile)
    docfile = os.path.join(input_folder, docfile)
    print(docfile)
    page_content = docx2python(docfile, html = True)
    print(page_content.text)
    soup = BeautifulSoup(page_content.text, "html.parser")
    for span in soup.find_all("span", attrs={"style": True}):
        span.unwrap()  # Removes the tag but keeps its content
    with open(outputfile, 'w', encoding='utf-8') as page_file:
        #print(page_file)
        page_file.write(str(soup))



# for filename in os.listdir(input_folder):
#     if filename.endswith('.docx'):
#         input_path = os.path.join(input_folder, filename)
#         output_filename = os.path.splitext(filename)[0] + ".html"
        
#         output_path = os.path.join(output_folder, output_filename)
        
#         doc = Document(input_path)
        
#         html_content = ''
#             # Convert each paragraph to a <p> tag
#         for para in doc.paragraphs:
#             html_content += f"<p>{para.text}</p>\n"
        
#         # print(html_content)
        
#         print(input_path)
      
#         with open(output_path, "w", encoding="utf-8") as html_file:
#             html_file.write(html_content)